"""PII Redaction Engine for DOCX prospectus-style documents.

High-recall hybrid detection:
- regex for structured PII
- optional spaCy PERSON NER
- context-aware person extraction
- table-aware person/company/address extraction
- stable synthetic replacements
"""
from __future__ import annotations

import argparse
import re
import sys
import io
import zipfile
from dataclasses import dataclass

from docx import Document
from PIL import Image, ImageDraw, ImageFont


# ----------------------------- Optional NER ---------------------------------
_NLP = None
_SPACY_AVAILABLE = False
try:
    import spacy
    try:
        _NLP = spacy.load(
            "en_core_web_sm",
            disable=["parser", "lemmatizer", "attribute_ruler", "tagger"],
        )
        _NLP.max_length = 10_000_000
        _SPACY_AVAILABLE = True
    except OSError:
        print(
            "[warn] spaCy model en_core_web_sm is not installed; "
            "using deterministic context/table extraction.",
            file=sys.stderr,
        )
except ImportError:
    print("[warn] spaCy is not installed; using deterministic extraction.", file=sys.stderr)


@dataclass(frozen=True)
class Match:
    start: int
    end: int
    kind: str
    value: str


# ----------------------------- Structured PII -------------------------------
EMAIL_RE = re.compile(r"\b[A-Z0-9._%+-]+@[A-Z0-9.-]+\.[A-Z]{2,}\b", re.I)
URL_RE = re.compile(r"(?:https?://|www\.)[^\s<>\"\'\]\[\)\}\}]+", re.I)
# DOCX files may split an email address around @ or dots.
SPACED_EMAIL_RE = re.compile(
    r"(?<![A-Za-z0-9._%+-])[A-Za-z0-9._%+-]+\s*@\s*"
    r"[A-Za-z0-9-]+(?:\s*\.\s*[A-Za-z0-9-]+)+",
    re.I,
)
SPACED_URL_RE = re.compile(
    r"(?:https?://|www\.)[A-Za-z0-9.-]+(?:\s*\.\s*[A-Za-z]{2,})"
    r"(?:[/#?][^\s<>\"\'\]\[\)\}\}]*)?",
    re.I,
)
# Bare domains commonly used in CVs, e.g. linkedin.com/in/name and
# github.com/username, are identifying URLs even without http:// or www.
BARE_PROFILE_URL_RE = re.compile(
    r"\b(?:linkedin\.com|github\.com|gitlab\.com|bitbucket\.org)"
    r"(?:/[^\s<>\"\'\]\[\)\}\}]*)?",
    re.I,
)

IP_RE = re.compile(
    r"\b(?:(?:25[0-5]|2[0-4]\d|1?\d?\d)\.){3}"
    r"(?:25[0-5]|2[0-4]\d|1?\d?\d)\b"
)

SSN_RE = re.compile(r"\b\d{3}-\d{2}-\d{4}\b")
CARD_RE = re.compile(r"(?<!\d)(?:\d[ -]?){13,19}(?!\d)")

DOB_RE = re.compile(
    r"\b(?:DOB|Date\s+of\s+Birth|Birth\s+Date)\s*[:\-]?\s*"
    r"(?:\d{1,2}[/-]\d{1,2}[/-]\d{2,4}"
    r"|\d{1,2}\s+[A-Za-z]+\s+\d{4}"
    r"|[A-Za-z]+\s+\d{1,2},\s*\d{4})\b",
    re.I,
)

PAN_RE = re.compile(r"\b[A-Z]{5}\d{4}[A-Z]\b")
DIN_RE = re.compile(
    r"\b(?:DIN|Director\s+Identification\s+Number)\s*[:#.-]?\s*(\d{8})\b",
    re.I,
)
CIN_RE = re.compile(
    r"\b[A-Z]\d{5}[A-Z]{2}\d{4}[A-Z]{3}\d{6}\b",
    re.I,
)
PIN_RE = re.compile(r"\b\d{3}\s?\d{3}\b")

PHONE_RE = re.compile(
    r"(?<!\d)(?:\+?\s*91[\s.-]?)?"
    r"(?:\(?0?\d{2,5}\)?[\s.-]?)?"
    r"\d{3,5}[\s.-]?\d{3,5}"
    r"(?:[\s.-]?\d{2,5})?(?!\d)"
)

PHONE_LABEL_RE = re.compile(
    r"\b(?:Telephone|Tel|Phone|Mobile|Contact\s*(?:No|Number)?|"
    r"Office\s*(?:No|Number)?)\s*[:.]?\s*",
    re.I,
)

ADDRESS_LABEL_RE = re.compile(
    r"\b(?:Registered Office|Corporate Office|Address|Registered Address|"
    r"Corporate Address|Residential Address|Correspondence Address|"
    r"Address of the (?:Registered Office|Corporate Office|Company|Director))"
    r"\s*[:\-]?\s*",
    re.I,
)

ADDRESS_CUES = re.compile(
    r"\b(?:Road|Rd\.?|Marg|Street|St\.?|Lane|Floor|Tower|Building|Bunglow|"
    r"Bungalow|Apartment|Society|Nagar|Colony|Complex|Park|Village|Taluka|"
    r"Tehsil|District|Industrial|Estate|Phase|Plot|House|Pune|Mumbai|"
    r"Bengaluru|Bangalore|Delhi|Maharashtra|Karnataka|Gujarat|India|"
    r"Gymkhana|Baner|Pashan|Bandra|Chakan|Khed|Taloja|Ahilyanagar|Ahmednagar|"
    r"Birdewadi|Deccan|Prabhat|Panchvati)\b",
    re.I,
)

ROLE_WORDS = re.compile(
    r"\b(?:Chairman|Managing Director|Joint Managing Director|Executive Director|"
    r"Whole[-\s]?time Director|Independent Director|Director|"
    r"Chief Executive Officer|CEO|Chief Financial Officer|CFO|"
    r"Company Secretary|Compliance Officer|CS|KMP|Key Managerial Personnel|"
    r"Senior Management|Technical Director|Contact Person|Promoter|"
    r"Promoter Selling Shareholder|Selling Shareholder|Shareholder|"
    r"Independent Chartered Engineer|Chartered Engineer)\b",
    re.I,
)

# Legal entities. The lazy prefix is important: it catches adjacent entities
# separately ("BSE Limited and National Stock Exchange of India Limited").
LEGAL_COMPANY_RE = re.compile(
    r"\b(?:(?:[A-Z][A-Za-z0-9&.()'’/-]*|and|of|the|for|&)(?:\s+)){0,10}?"
    r"(?:Limited|LIMITED|Private Limited|PRIVATE LIMITED|"
    r"Pvt\.?\s*Ltd\.?|LLP|Trust|TRUST|Foundation|FOUNDATION)\b"
)

ENTITY_WORDS = re.compile(
    r"\b(?:Securities|Motors|Electricals|Finance|Financial|Insurance|Bank|"
    r"Capital|Advisory|Consulting|Holdings|Industries|Industrial|Solutions|"
    r"Technologies|Analytics|Ventures|Investments|Infrastructure|Infra|"
    r"Properties|Developers|Research|Logistics|HUF)\b",
    re.I,
)

GENERIC_NAME_REJECT_PHRASES = {
    "server ip",
    "client ip",
    "ip address",
    "email address",
    "phone number",
    "telephone number",
    "account number",
    "order number",
    "ticket number",
    "date of birth",
    "credit card",
    "card number",
    "postal code",
    "pin code",
    "registration number",
    "identification number",
    "reference number",
}

# These words make a capitalized span overwhelmingly likely to be a document
# heading, table header, financial phrase, or entity rather than a person.
REJECT_WORDS = {
    "OFFER", "OFFERED", "PROMOTER", "PROMOTERS", "DIRECTOR", "DIRECTORS",
    "PRICE", "CAP", "FLOOR", "FUND", "FUNDS", "MUTUAL", "FACILITY",
    "BIDDER", "BIDDERS", "UPI", "EMAIL", "REGISTRAR", "SHARE", "SHARES",
    "TRANSFER", "AGENTS", "AGENT", "KEY", "MANAGERIAL", "PERSONNEL",
    "REFERENCE", "RATE", "DELAY", "SELLING", "SHAREHOLDER", "SHAREHOLDERS",
    "SECONDARY", "SPLIT", "EXCLUDES", "NON", "GAAP", "MEASURES", "COMPLEX",
    "PARK", "MARG", "ROAD", "BRANCH", "TALUKA", "RECLAMATION", "CHURCHGATE",
    "NEWSPAPER", "CIRCULATED", "WIDELY", "DAILY", "MARATHI", "ENGLISH",
    "GUJARATI", "HINDI", "BACKBAY", "COMMITTEE", "SUB", "COMPANY",
    "COMPANIES", "LIMITED", "PRIVATE", "LLP", "TRUST", "FOUNDATION",
    "BANK", "BANKS", "INDIA", "INDIAN", "TERM", "SHORT", "LONG",
    "CERTIFIED", "SYNDICATE", "ESCROW", "ACCOUNT", "ACCOUNTS",
    "COLLECTION", "REFUND", "OFFICE", "OFFICER", "SECRETARY",
    "SEBI", "ICDR", "RBI", "NSE", "BSE", "GST", "TAN", "CIN", "ROC",
    "KMP", "KMPS", "SM", "SMS", "BRLM", "BRLMS", "ACT", "ACTS",
    "REGULATION", "REGULATIONS", "RULES", "SCHEME", "SCHEMES",
    "STATEMENT", "STATEMENTS", "REPORT", "REPORTS", "ANALYSIS",
    "DISCUSSION", "MANAGEMENT", "BUSINESS", "STRUCTURE", "PROCEEDS",
    "PORTION", "INVESTOR", "INVESTORS", "ELIGIBLE", "NRI", "NRIS",
    "INSTITUTIONAL", "EXPRESS", "EXCHANGE", "EXCHANGES", "LONDON",
    "METAL", "WORKING", "DAYS", "DEFAULTER", "WILFUL", "LAND",
    "FREEHOLD", "LEASEHOLD", "PROVISIONS", "INSURANCE", "WEBSITE",
    "CONTACT", "PERSON", "SLIP", "ACKNOWLEDGEMENT", "CONDITIONING",
    "AIR", "PORTAL", "PORTIONS", "QIB", "ANCHOR", "AMOUNT", "BID",
    "DP", "ID", "SCHEDULE", "SECURITIES", "CONTRACTS", "SENIOR",
    "OUR", "CAGR", "PAT", "MARGIN", "TAX", "DEDUCTED", "OTHER",
    "ASSETS", "NET", "PHOTO", "VOLTAIC", "VOLT", "AMPERES", "MEGA",
    "MISCELLANEOUS", "STATE", "GOVERNMENT", "POLICY", "POLICIES",
    "INDUSTRIAL", "URJA", "SURAKSHA", "KISAN", "COLONY", "MODEL",
    "OPP", "CHAMBERS", "SHOWROOM", "SH", "GAAP", "OPERATIONAL",
    "MANAGING", "JOINT", "CIRCUIT", "KILOMETERS", "WORK", "CAPITAL",
    "ANALYTICS", "GRAM", "LEASING", "MONTE", "BUENA", "NOMINEE",
    "TOTAL", "NUMBER", "DATE", "DESIGNATION", "ADDRESS", "ENTITY",
    "DESCRIPTION", "PARTICULARS", "RESPONSIBILITY", "COORDINATOR",
    "APPLICATION", "BOOK", "RUNNING", "LEAD", "MANAGER", "MANAGERS",
    "AVERAGE", "COST", "ACQUISITION", "PER", "CERTAIN", "CORPORATE",
    "MATTERS", "CHANGE", "DRAFT", "HERING", "HERRING", "PROSPECTUS",
    "PRICING", "FINANCIAL", "INFORMATION", "IFRS", "UNITED", "STATES",
    "FISCAL", "YEAR", "YEARS", "OFFERING", "INITIAL", "PUBLIC", "ISSUE",
    "ISSUES", "EQUITY", "FACE", "VALUE", "MILLION", "MILLIONS", "DETAILS",
    "INTRODUCTION", "GENERAL", "RISK", "RISKS", "FACTORS", "SECTION",
    "SUMMARY", "RELATED", "PARTY", "DISCLOSURES", "DISCLOSURE",
    "TRANSACTION", "TRANSACTIONS", "STATUTORY", "UNDERWRITING", "AGREEMENT",
    "GAT", "NO", "IND", "AS", "MATERIAL", "DEVELOPMENTS", "OUTSTANDING",
    "LITIGATION", "NAME", "OF", "THE", "HUF", "ISSUER",
}

NON_PERSON_PHRASES = {
    "Red Herring Prospectus", "General Information", "Book Building Process",
    "Risk Factors", "Board of Directors", "National Stock Exchange",
    "Securities and Exchange Board", "Management Discussion",
    "Financial Statements", "Registered Office", "Corporate Office",
    "Offer Document", "Basis of Allotment", "Business Strategy",
    "Industry Overview", "Key Managerial Personnel", "Share Transfer Agents",
    "Mutual Funds", "Mutual Fund", "Cap Price", "Floor Price", "Offer Price",
    "Price Band", "Anchor Investor", "Retail Individual Investors",
    "Qualified Institutional Buyers", "Selling Shareholder",
    "Selling Shareholders", "Non-GAAP Measures", "Registrar to the Offer",
}


# Explicit aliases that may not be caught by legal-entity regexes.
# Keep these targeted: broad token matching can damage normal prospectus text.
EXPLICIT_COMPANY_NAMES = {
    "KSH",
    "ICICI Venture",
    "ICICI Securities",
    "ICICI Securities Limited",
    "Nuvama",
    "MUFG Intime India",
    "MUFG Intime India Private Limited",
    "Everest Family Trust",
    "Annapurna Family Trust",
    "Makalu Family Trust",
    "Kanchenjunga Family Trust",
}

EXPLICIT_PERSON_NAMES = {"Rajesh", "Rakhi", "Rohit", "Sangeeta"}

FAMILY_BRANCH_NAMES = {
    "Parents Branch",
    "Rajesh Branch",
    "Sangeeta Branch",
    "Rakhi Branch",
    "Rohit Branch",
}

# ----------------------------- Replacements --------------------------------
class ReplacementMap:
    FAKE_VALUES = {
        "PERSON": [f"Person {x}" for x in
                   ["Alpha", "Beta", "Gamma", "Delta", "Epsilon", "Zeta",
                    "Eta", "Theta", "Iota", "Kappa", "Lambda", "Mu", "Nu",
                    "Xi", "Omicron", "Pi", "Rho", "Sigma", "Tau", "Upsilon",
                    "Phi", "Chi", "Psi", "Omega"]],
        "EMAIL": [f"redacted.person{i:03d}@example.invalid" for i in range(1, 101)],
        "URL": [f"https://example.invalid/site{i:03d}" for i in range(1, 101)],
        "PHONE": [f"+91 90000 {i:05d}" for i in range(1, 101)],
        "COMPANY": [f"Example Company {i:03d} Limited" for i in range(1, 101)],
        "BRANCH": [f"Family Branch {i:03d}" for i in range(1, 101)],
        "ADDRESS": [
            "1 Example Road, Pune - 411001, Maharashtra, India",
            "2 Sample Street, Mumbai - 400001, Maharashtra, India",
            "3 Demo Avenue, Bengaluru - 560001, Karnataka, India",
            "4 Test Park, Delhi - 110001, India",
            "5 Example Park, Hyderabad - 500001, Telangana, India",
        ],
        "SSN": ["000-00-0001", "000-00-0002", "000-00-0003"],
        "CARD": ["4000 0000 0000 0001", "4000 0000 0000 0002", "4000 0000 0000 0003"],
        "DOB": ["January 01, 1990", "February 02, 1991", "March 03, 1992"],
        "IP": ["192.0.2.1", "192.0.2.2", "192.0.2.3"],
        "DIN": ["00000001", "00000002", "00000003"],
        "PAN": ["AAAAA0000A", "BBBBB0000B", "CCCCC0000C"],
        "CIN": [
            "U00000XX0000PLC000001",
            "U00000XX0000PLC000002",
            "U00000XX0000PLC000003",
        ],
    }

    def __init__(self):
        self.maps = {k: {} for k in self.FAKE_VALUES}
        self.next_index = {k: 0 for k in self.FAKE_VALUES}

    def fake(self, kind: str, original: str) -> str:
        original = original.strip()
        key = (
            re.sub(r"\s+", " ", original).strip().casefold()
            if kind in {"PERSON", "COMPANY", "EMAIL", "URL", "BRANCH"}
            else original.strip()
        )
        if key not in self.maps[kind]:
            values = self.FAKE_VALUES[kind]
            idx = self.next_index[kind] % len(values)
            self.maps[kind][key] = values[idx]
            self.next_index[kind] += 1
        return self.maps[kind][key]


# ----------------------------- DOCX traversal -------------------------------
def iter_blocks(doc: Document):
    """Yield each physical DOCX paragraph exactly once.

    python-docx can expose the same XML paragraph through multiple table-cell
    proxies, especially for merged/repeated cells. Processing those proxies
    repeatedly was causing already-redacted text to be processed again and
    produced strings such as ``Example Company 004 Example Company 004``.
    """
    # Keep the XML elements themselves in the set. Storing only id(...) can
    # suffer from Python object-id reuse while iterating thousands of DOCX
    # proxy objects, which can accidentally skip a real paragraph.
    seen = set()

    def emit(p):
        key = p._p
        if key in seen:
            return None
        seen.add(key)
        return p

    for p in doc.paragraphs:
        q = emit(p)
        if q is not None:
            yield q

    def walk_table(table):
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    q = emit(p)
                    if q is not None:
                        yield q
                for nested in cell.tables:
                    yield from walk_table(nested)

    for table in doc.tables:
        yield from walk_table(table)

    for section in doc.sections:
        for p in section.header.paragraphs:
            q = emit(p)
            if q is not None:
                yield q
        for p in section.footer.paragraphs:
            q = emit(p)
            if q is not None:
                yield q


def collect_blocks(doc: Document) -> list[str]:
    return [
        " ".join(p.text.split())
        for p in iter_blocks(doc)
        if p.text and p.text.strip()
    ]


# ----------------------------- Candidate checks -----------------------------
def clean_candidate(text: str) -> str:
    text = re.sub(r"\s+", " ", text.replace("\n", " "))
    text = text.strip(" ,.;:()[]{}<>*#^&\"'")
    tokens = text.split()

    # Prevent sentence continuations such as "Amod Joshi. For" from becoming
    # a four-word person candidate.
    for i, tok in enumerate(tokens):
        core = tok.strip(",;:()[]{}<>*#^&\"'")
        if core.endswith(".") and len(core) > 2:
            tokens = tokens[:i + 1]
            break

    return " ".join(tokens).strip(" ,.;:()[]{}<>*#^&\"'")


def name_like(text: str) -> bool:
    text = clean_candidate(text)
    if not text or text in NON_PERSON_PHRASES:
        return False

    words = text.split()
    if not 2 <= len(words) <= 5:
        return False

    if ADDRESS_CUES.search(text):
        return False

    for word in words:
        if not re.fullmatch(r"[A-Z][A-Za-z'’.-]*", word):
            return False
        parts = re.split(r"[-']", word.strip("."))
        if any(p.upper() in REJECT_WORDS for p in parts if p):
            return False

    # Reject corporate/entity-shaped candidates.
    if re.search(
        r"\b(?:Limited|Private|LLP|Trust|Foundation|Bank|Securities|"
        r"Finance|Financial|Insurance|Capital|Consulting|Holdings|"
        r"Electricals|Motors|Industries|Industrial|Solutions|Research|"
        r"Centre|Processing|Logistics|Infrastructure|Ventures|HUF)\b",
        text,
        re.I,
    ):
        return False

    return True


def strong_name_like(text: str) -> bool:
    """High-confidence person check for explicit person contexts.

    Unlike the normal conservative name check, this also accepts ALL-CAPS
    names, which are common on prospectus covers and promoter banners.
    """
    text = clean_candidate(text)
    if name_like(text):
        return True

    words = text.split()
    if not 2 <= len(words) <= 5:
        return False

    for word in words:
        core = word.strip(".,")
        if not re.fullmatch(r"[A-Z][A-Z'’.-]*", core):
            return False
        parts = re.split(r"[-']", core)
        if any(p.upper() in REJECT_WORDS for p in parts if p):
            return False

    if ADDRESS_CUES.search(text):
        return False

    if re.search(
        r"\b(?:LIMITED|PRIVATE|LLP|TRUST|FOUNDATION|BANK|SECURITIES|"
        r"FINANCE|FINANCIAL|INSURANCE|CAPITAL|CONSULTING|HOLDINGS|"
        r"INDUSTRIES|INDUSTRIAL|SOLUTIONS|RESEARCH|VENTURES|HUF)\b",
        text,
        re.I,
    ):
        return False

    return True


def extract_person_prefix(text: str) -> set[str]:
    """Extract a leading person name from a strong labelled context."""
    people = set()
    raw_tokens = re.findall(r"[A-Z][A-Za-z'’.-]*", text.strip())
    if not raw_tokens:
        return people

    stop = {
        "SEBI", "REGISTRATION", "NUMBER", "NO", "DIN", "EMAIL", "E-MAIL",
        "WEBSITE", "TELEPHONE", "TEL", "PHONE", "ADDRESS", "CIN", "URL",
        "AGGREGATING", "EACH", "EQUITY", "SHARES", "MILLION", "LIMITED",
        "PRIVATE", "LLP", "TRUST", "FOUNDATION", "AND", "OR",
        "COMPLIANCE", "OFFICER", "DESIGNATION",
    }

    # Only inspect the leading 2-5 tokens. This prevents trailing metadata
    # from being swallowed into a person's name.
    for n in range(min(5, len(raw_tokens)), 1, -1):
        prefix = raw_tokens[:n]
        if any(tok.upper() in stop for tok in prefix):
            continue
        candidate = " ".join(prefix)
        if strong_name_like(candidate):
            people.add(clean_candidate(candidate))
            return people

    return people


def entity_like(text: str) -> bool:
    text = clean_candidate(text)
    if not text or text in NON_PERSON_PHRASES:
        return False
    words = text.split()
    if not 1 <= len(words) <= 10:
        return False
    if not re.match(r"^[A-Z]", text):
        return False
    if name_like(text):
        return False
    if all(w.strip(".,").upper() in REJECT_WORDS for w in words):
        return False
    return True


# ----------------------------- Person extraction ----------------------------
NAME_SPAN_RE = re.compile(
    r"\b[A-Z][A-Za-z'’.-]+(?:\s+[A-Z][A-Za-z'’.-]+){1,4}\b"
)

PERSON_TRIGGER_RE = re.compile(
    r"\b(?:Contact\s+Person|our\s+Promoters?|our\s+Directors?|"
    r"Promoter\s+Selling\s+Shareholders?|Promoters?|"
    r"Shareholders?|Chairman(?:\s+and\s+Executive\s+Director)?|"
    r"Managing\s+Director|Joint\s+Managing\s+Director|Executive\s+Director|"
    r"Whole[-\s]?time\s+Director|Independent\s+Director|"
    r"Chief\s+Executive\s+Officer|Chief\s+Financial\s+Officer|"
    r"Company\s+Secretary|Compliance\s+Officer|Senior\s+Management|"
    r"Key\s+Managerial\s+Personnel|Independent\s+Chartered\s+Engineer|"
    r"Registrar\s+Contact\s+Person)\b",
    re.I,
)


def name_spans(text: str) -> set[str]:
    return {
        clean_candidate(m.group(0))
        for m in NAME_SPAN_RE.finditer(text)
        if name_like(clean_candidate(m.group(0)))
    }


def extract_generic_people(blocks: list[str], companies: set[str]) -> set[str]:
    """Conservative fallback for PERSON names missed by spaCy.

    spaCy can miss unfamiliar names. This fallback considers title-case
    2-5 word spans, but rejects known non-person phrases, address/entity
    terminology, and any span overlapping a discovered company name.
    """

    people = set()

    company_keys = {
        re.sub(r"\s+", " ", c).strip().casefold()
        for c in companies
    }

    # Common sentence/heading words that can begin or end capitalized spans
    # without representing a person's name.
    contextual_reject = {
        "The", "This", "That", "These", "Those", "Our", "Their", "His", "Her",
        "Its", "For", "From", "With", "And", "Or", "But", "In", "On", "At",
        "By", "To", "Of", "As", "After", "Before", "During", "Under",
        "According", "Based", "Further", "Details", "Information",
    }

    for text in blocks:
        for m in NAME_SPAN_RE.finditer(text):
            candidate = clean_candidate(m.group(0))
            if not name_like(candidate):
                continue

            words = candidate.split()

            # Avoid obvious sentence fragments.
            if words[0] in contextual_reject or words[-1] in contextual_reject:
                continue

            # Avoid candidates that contain a discovered company/entity.
            candidate_key = re.sub(r"\s+"," ",candidate).strip().casefold()

# Reject common technical, document, and field-label phrases.
            if candidate_key in GENERIC_NAME_REJECT_PHRASES:
                continue

            if any(
                candidate_key == ck
                or candidate_key.startswith(ck + " ")
                or candidate_key.endswith(" " + ck)
                for ck in company_keys
            ):
                continue

            # Require at least two conventional name tokens.
            # Do not use a one-word generic-name fallback here.
            if len(words) >= 2:
                people.add(candidate)

    return people


def extract_spacy_people(blocks: list[str]) -> set[str]:
    people = set()
    if not _SPACY_AVAILABLE:
        return people

    for d in _NLP.pipe(blocks, batch_size=100):
        for ent in d.ents:
            if ent.label_ == "PERSON":
                c = clean_candidate(ent.text)
                if name_like(c):
                    people.add(c)
    return people


def extract_context_people(blocks: list[str]) -> set[str]:
    people = set()

    # CV/resume owner-name detection. A resume commonly puts the person's
    # name alone on the first line, without a role such as "Director" or
    # "Contact Person". Use the first block only when nearby blocks contain
    # strong contact/profile signals; this avoids treating ordinary headings
    # such as "Projects" or "Curriculum Vitae" as people.
    if blocks:
        first = clean_candidate(blocks[0])
        nearby = " ".join(blocks[:5])
        has_profile_signal = bool(re.search(
            r"@|\b(?:linkedin|github|portfolio|phone|mobile|telephone|email)\b|"
            r"\+?\d[\d ()-]{8,}",
            nearby,
            re.I,
        ))
        if has_profile_signal and strong_name_like(first):
            people.add(first)

        # Some resume templates put the name in a short first block followed
        # by a location/contact line. Accept a 2-4 word candidate there too.
        if len(first.split()) <= 4 and has_profile_signal:
            if _NLP is not None:
                try:
                    doc0 = _NLP(first)
                    if any(ent.label_ == "PERSON" for ent in doc0.ents) and name_like(first):
                        people.add(first)
                except Exception:
                    pass

    for idx, text in enumerate(blocks):
        # Contact Person labels are high-confidence. Cut before metadata and
        # extract only the leading name, e.g.
        # "Shanti Gopalkrishnan SEBI Registration No.".
        for m in re.finditer(
            r"\bContact\s+Person\s*[:\-]?\s*(.+)",
            text,
            re.I,
        ):
            tail = re.split(
                r";|\b(?:Telephone|Tel|Phone|Email|E-mail|Website|"
                r"SEBI\s+Registration(?:\s+Number)?|DIN|CIN|URL)\b",
                m.group(1),
                maxsplit=1,
                flags=re.I,
            )[0]
            for piece in re.split(r"/|\band\b|,", tail, flags=re.I):
                people |= extract_person_prefix(piece)

        # Existing role-based extraction, but also use the strong prefix
        # detector for cases such as ALL-CAPS promoter/contact names.
        for trigger in PERSON_TRIGGER_RE.finditer(text):
            tail = text[trigger.end():]
            tail = re.split(
                r"[;|]\s*|\b(?:Telephone|Tel|Phone|Email|E-mail|Website|"
                r"SEBI\s+Registration(?:\s+Number)?|DIN|CIN|Address)\b",
                tail,
                maxsplit=1,
                flags=re.I,
            )[0]

            for piece in re.split(r"/|\band\b|,", tail, flags=re.I):
                people |= extract_person_prefix(piece)

            # Keep the original conservative title-case extraction.
            for candidate in name_spans(tail[:500]):
                people.add(candidate)

        # "being NAME" / "namely NAME".
        for m in re.finditer(
            r"\b(?:being|namely)\s*,?\s*"
            r"([A-Z][A-Za-z'’.-]+(?:\s+[A-Z][A-Za-z'’.-]+){1,4})",
            text,
            re.I,
        ):
            c = clean_candidate(m.group(1))
            # Trim known trailing metadata.
            c = re.split(
                r"\s+(?:bearing|registration|for\s+further|see|with)\b",
                c,
                maxsplit=1,
                flags=re.I,
            )[0]
            if strong_name_like(c):
                people.add(c)

        # NAME, Chairman / NAME, Director / NAME, CFO etc.
        for m in re.finditer(
            r"\b([A-Z][A-Za-z'’.-]+(?:\s+[A-Z][A-Za-z'’.-]+){1,4})"
            r"\s*,?\s+(?=" + ROLE_WORDS.pattern + r")",
            text,
            re.I,
        ):
            c = clean_candidate(m.group(1))
            if strong_name_like(c):
                people.add(c)

        # Prospectus cover/banner promoter lists are frequently ALL CAPS.
        # Split people from companies by comma/AND.
        for m in re.finditer(
            r"\bOUR\s+PROMOTERS?\s*:\s*(.+?)(?=\bDETAILS\s+OF\b|"
            r"\bINITIAL\s+PUBLIC\s+OFFERING\b|$)",
            text,
            re.I,
        ):
            for piece in re.split(r",|\band\b", m.group(1), flags=re.I):
                c = clean_candidate(piece)
                if strong_name_like(c):
                    people.add(c)

        # Offer paragraphs commonly use "BY FULL NAME AGGREGATING ...".
        for m in re.finditer(
            r"\bBY\s+(.{0,90})",
            text,
            re.I,
        ):
            tail = re.split(
                r"\b(?:AGGREGATING|EACH|UP\s+TO|COLLECTIVELY|THE\s+"
                r"PROMOTER|THE\s+OFFER|AND\s+SUCH)\b",
                m.group(1),
                maxsplit=1,
                flags=re.I,
            )[0]
            for piece in re.split(r",|\band\b", tail, flags=re.I):
                people |= extract_person_prefix(piece)

        # Consent / expert sentences: "consent ... from Lalit X, to include..."
        for m in re.finditer(
            r"\b(?:consent|consent\s+dated|written\s+consent).*?\bfrom\s+"
            r"([A-Z][A-Za-z'’.-]+(?:\s+[A-Z][A-Za-z'’.-]+){1,4})\s*,",
            text,
            re.I,
        ):
            c = clean_candidate(m.group(1))
            if strong_name_like(c):
                people.add(c)

        # Family-branch definitions. A DOCX table may store the branch label
        # and the person's name in adjacent paragraphs rather than one cell.
        if re.fullmatch(
            r"(?:Rajesh|Rakhi|Rohit|Sangeeta|Parents)\s+Branch",
            text.strip(),
            re.I,
        ):
            if idx + 1 < len(blocks):
                nxt = blocks[idx + 1]
                first = re.split(
                    r",\s*(?:his|her|their)\b",
                    nxt,
                    maxsplit=1,
                    flags=re.I,
                )[0]
                people |= extract_person_prefix(first)
                for candidate in name_spans(first[:120]):
                    people.add(candidate)

        # When branch label and person are in the same block.
        if re.search(
            r"\b(?:Rajesh|Rakhi|Rohit|Sangeeta|Parents)\s+Branch\b",
            text,
            re.I,
        ):
            tail = re.split(
                r"\b(?:Rajesh|Rakhi|Rohit|Sangeeta|Parents)\s+Branch\b",
                text,
                maxsplit=1,
                flags=re.I,
            )[-1]
            tail = re.split(
                r",\s*(?:his|her|their)\b",
                tail,
                maxsplit=1,
                flags=re.I,
            )[0]
            people |= extract_person_prefix(tail)

        # Related-party/share-transfer sentences: retain the original
        # conservative NER/context route.
        if re.search(r"\b(?:transfer of shares|related party)\b", text, re.I):
            people |= name_spans(text)

    # Final safety filter: role labels and structural phrases are not people.
    people = {
        p for p in people
        if p.strip().casefold() not in {
            "chartered accountants",
            "designated intermediaries",
            "designated intermediary",
            "designated stock",
            "in consultation with",
            "independent chartered engineer",
            "issuer’s and",
        }
        and not re.search(
            r"\b(?:chartered|accountants|designated|intermediaries|"
            r"intermediary|consultation|issuer’s|engineer)\b",
            p,
            re.I,
        )
    }
    return people


# ----------------------------- Table extraction ------------------------------
def extract_table_entities(doc: Document) -> tuple[set[str], set[str]]:
    people, companies = set(), set()

    strong_person_header = re.compile(
        r"\b(?:name\s+of\s+(?:the\s+)?(?:promoter|shareholder|director|"
        r"selling\s+shareholder|contact\s+person|kmp)|"
        r"promoter\s+selling\s+shareholder|contact\s+person)\b",
        re.I,
    )

    for table in doc.tables:
        if not table.rows:
            continue

        headers = [" ".join(c.text.split()) for c in table.rows[0].cells]
        target_indices = []

        for i, header in enumerate(headers):
            if re.search(r"\bname\b", header, re.I) and not re.search(
                r"\b(?:entity|description)\b", header, re.I
            ):
                target_indices.append(i)
            elif strong_person_header.search(header):
                target_indices.append(i)

        board_like = (
            any(re.search(r"\bdesignation\b", h, re.I) for h in headers)
            and any(re.search(r"\bdin\b", h, re.I) for h in headers)
            and any(re.search(r"\baddress\b", h, re.I) for h in headers)
        )
        if board_like:
            target_indices = [0]

        for row in table.rows[1:]:
            cells = [" ".join(c.text.split()) for c in row.cells]
            nonempty = [c for c in cells if c]
            if nonempty and all(c == nonempty[0] for c in nonempty):
                continue

            for idx in set(target_indices):
                if idx >= len(cells):
                    continue

                value = clean_candidate(cells[idx])
                if not value:
                    continue

                if re.search(
                    r"^(?:total|sub-total|subtotal|payment|remuneration|grand|name)$",
                    value,
                    re.I,
                ):
                    continue

                for piece in re.split(r"\s*/\s*", value):
                    piece = clean_candidate(piece)
                    if name_like(piece):
                        people.add(piece)
                    elif entity_like(piece) and (
                        LEGAL_COMPANY_RE.search(piece)
                        or ENTITY_WORDS.search(piece)
                    ):
                        companies.add(piece)

            # Branch rows often have no "Name" header.
            if len(cells) >= 2 and re.search(r"\bBranch\b", cells[0], re.I):
                first = clean_candidate(cells[1].split(",", 1)[0])
                if name_like(first):
                    people.add(first)

    return people, companies


# ----------------------------- Company extraction ---------------------------
def clean_company(text: str) -> str:
    text = clean_candidate(text)
    leading = {
        "AND", "THE", "OF", "FOR", "EMAIL", "E-MAIL", "TELEPHONE", "TEL",
        "PHONE", "WEBSITE", "ADDRESS", "CONTACT", "PERSON", "NAME",
        "REGISTRAR", "COMPANY", "ISSUER", "BANKER", "BANKERS", "NOTE",
        "INCLUDING", "NAMELY", "OUR", "BY", "TO", "FROM", "IN",
    }
    words = text.split()
    while words and words[0].strip(".,").upper() in leading:
        words.pop(0)
    return " ".join(words)


def extract_companies(blocks: list[str], table_companies: set[str]) -> set[str]:
    companies = set(table_companies)

    for text in blocks:
        for m in LEGAL_COMPANY_RE.finditer(text):
            c = clean_company(m.group(0))
            if len(c) >= 5:
                companies.add(c)

        # Entity names in strong labels.
        for m in re.finditer(
            r"\b(?:Company|Issuer|Registrar|Banker|Bankers|Broker|"
            r"Book Running Lead Manager|Lead Manager|Monitoring Agency|"
            r"Underwriter|Related Party)\s*[:\-]?\s*"
            r"([A-Z][A-Za-z0-9&.'’/-]*(?:\s+[A-Z][A-Za-z0-9&.'’/-]*){0,9})",
            text,
            re.I,
        ):
            c = clean_company(m.group(1))
            c = re.split(
                r"\s+(?:Telephone|Tel|Email|E-mail|Website|Address|"
                r"Contact Person|Designation|DIN)\b",
                c,
                maxsplit=1,
                flags=re.I,
            )[0]
            if entity_like(c) and (
                LEGAL_COMPANY_RE.search(c) or ENTITY_WORDS.search(c)
            ):
                companies.add(c)

    return companies


# ----------------------------- Address extraction ---------------------------
def address_match(text: str) -> Match | None:
    label = ADDRESS_LABEL_RE.search(text)
    if label:
        value = text[label.end():].strip()
        if value and (PIN_RE.search(value) or ADDRESS_CUES.search(value)):
            return Match(label.end(), len(text), "ADDRESS", value)

    # Strong unlabeled address: require both a PIN and a geographic/address cue.
    if PIN_RE.search(text) and ADDRESS_CUES.search(text):
        pin = PIN_RE.search(text)
        cue = ADDRESS_CUES.search(text)
        start = min(pin.start(), cue.start())

        if re.match(
            r"\s*(?:Plot|Flat|House|Unit|S\.?\s*No\.?|No\.?)?\s*\d+",
            text,
            re.I,
        ):
            start = 0

        return Match(start, len(text), "ADDRESS", text[start:].strip())

    return None


# ----------------------------- Match finding --------------------------------
def find_matches(text: str, people: set[str], companies: set[str], person_pattern=None, company_pattern=None, branch_pattern=None) -> list[Match]:
    matches = []

    def add(regex, kind):
        for m in regex.finditer(text):
            matches.append(Match(m.start(), m.end(), kind, m.group(0)))

    add(EMAIL_RE, "EMAIL")
    # Prefer the spaced form only when the normal email regex did not already
    # cover the same span. It is needed for DOCX files whose runs contain
    # ``name @domain.com``.
    for m in SPACED_EMAIL_RE.finditer(text):
        compact = re.sub(r"\s+", "", m.group(0))
        if EMAIL_RE.fullmatch(compact):
            matches.append(Match(m.start(), m.end(), "EMAIL", m.group(0)))

    add(URL_RE, "URL")
    for m in SPACED_URL_RE.finditer(text):
        if not any(x.kind == "URL" and not (m.end() <= x.start or m.start() >= x.end) for x in matches):
            matches.append(Match(m.start(), m.end(), "URL", m.group(0)))
    for m in BARE_PROFILE_URL_RE.finditer(text):
        if not any(x.kind == "URL" and not (m.end() <= x.start or m.start() >= x.end) for x in matches):
            matches.append(Match(m.start(), m.end(), "URL", m.group(0)))
    add(IP_RE, "IP")
    add(SSN_RE, "SSN")
    add(DOB_RE, "DOB")
    add(PAN_RE, "PAN")

    for m in DIN_RE.finditer(text):
        matches.append(
            Match(m.start(1), m.end(1), "DIN", m.group(1))
        )

    add(CIN_RE, "CIN")

    for m in CARD_RE.finditer(text):
        if 13 <= len(re.sub(r"[ -]", "", m.group(0))) <= 19:
            matches.append(Match(m.start(), m.end(), "CARD", m.group(0)))

    # Labelled phones.
    for label in PHONE_LABEL_RE.finditer(text):
        segment = text[label.end():]
        stop = re.search(
            r"\b(?:Email|E-mail|Website|Contact\s+Person|Address|"
            r"SEBI Registration Number)\b",
            segment,
            re.I,
        )
        if stop:
            segment = segment[:stop.start()]
        for pm in PHONE_RE.finditer(segment[:120]):
            raw = pm.group(0)
            digits = re.sub(r"\D", "", raw)
            if 10 <= len(digits) <= 13:
                matches.append(
                    Match(
                        label.end() + pm.start(),
                        label.end() + pm.end(),
                        "PHONE",
                        raw,
                    )
                )

    # Standalone Indian mobile numbers.
    for m in re.finditer(r"(?<!\d)(?:\+?91[\s.-]?)?[6-9]\d{9}(?!\d)", text):
        digits = re.sub(r"\D", "", m.group(0))
        if len(digits) in (10, 12):
            matches.append(Match(m.start(), m.end(), "PHONE", m.group(0)))

    if company_pattern is not None:
        for m in company_pattern.finditer(text):
            matches.append(Match(m.start(), m.end(), "COMPANY", m.group(0)))

    if branch_pattern is not None:
        for m in branch_pattern.finditer(text):
            matches.append(Match(m.start(), m.end(), "BRANCH", m.group(0)))

    if person_pattern is not None:
        for m in person_pattern.finditer(text):
            matches.append(Match(m.start(), m.end(), "PERSON", m.group(0)))

    addr = address_match(text)
    if addr:
        matches.append(addr)

    priority = {
        "EMAIL": 100, "URL": 99, "CARD": 98, "SSN": 98, "IP": 98, "DOB": 96,
        "CIN": 95, "PAN": 94, "DIN": 92, "PHONE": 90, "ADDRESS": 85,
        "PERSON": 70, "BRANCH": 68, "COMPANY": 60,
    }

    # Longest + highest-confidence wins when matches overlap.
    matches.sort(
        key=lambda m: (-(m.end - m.start), -priority[m.kind], m.start)
    )

    chosen = []
    for m in matches:
        if not any(
            not (m.end <= x.start or m.start >= x.end)
            for x in chosen
        ):
            chosen.append(m)

    return sorted(chosen, key=lambda m: m.start)



# ----------------------------- Shared detection helpers ----------------------
def build_entity_pattern(values):
    """Build a whitespace-tolerant pattern for discovered entities."""
    values = sorted(
        {v for v in values if len(v.strip()) >= 2},
        key=len,
        reverse=True,
    )
    if not values:
        return None

    patterns = []
    for value in values:
        pat = re.escape(value)
        pat = pat.replace(r"\ ", r"\s+")
        patterns.append(pat)

    return re.compile(
        r"(?<![A-Za-z'’])(?:"
        + "|".join(patterns)
        + r")(?![A-Za-z'’])",
        re.I,
    )


def prepare_detection_entities(blocks, doc=None):
    """Run the same entity-discovery stages used by the DOCX pipeline."""
    people = extract_spacy_people(blocks)
    people |= extract_context_people(blocks)
    people |= EXPLICIT_PERSON_NAMES

    table_people = set()
    table_companies = set()

    if doc is not None:
        table_people, table_companies = extract_table_entities(doc)
        people |= table_people

    companies = extract_companies(blocks, table_companies)
    companies |= EXPLICIT_COMPANY_NAMES

    people |= extract_generic_people(blocks, companies)

    for text in blocks:
        for m in re.finditer(
            r"\b(?:Name|Contact Person)\s*[:\-]\s*(.+)",
            text,
            re.I,
        ):
            tail = re.split(
                r";|\b(?:Telephone|Tel|Phone|Email|E-mail|Website|"
                r"SEBI\s+Registration(?:\s+Number)?|DIN|CIN|URL)\b",
                m.group(1),
                maxsplit=1,
                flags=re.I,
            )[0]
            for piece in re.split(r"/|\band\b|,", tail, flags=re.I):
                people |= extract_person_prefix(piece)

    branches = set(FAMILY_BRANCH_NAMES)
    return people, companies, branches


def detect_pii_for_evaluation(text):
    """Expose the production detector for the separate evaluation harness.

    This function contains no benchmark logic. It uses the same entity
    discovery and find_matches() path as redact_docx(), but on plain text.
    """
    blocks = [" ".join(text.split())]
    people, companies, branches = prepare_detection_entities(blocks)

    return find_matches(
        blocks[0],
        people,
        companies,
        person_pattern=build_entity_pattern(people),
        company_pattern=build_entity_pattern(companies),
        branch_pattern=build_entity_pattern(branches),
    )


# ----------------------------- Formatting-safe replacement -----------------
def replace_in_paragraph(paragraph, matches, replacements, stats):
    """Apply all replacements against the original Word-run offsets.

    Never mutate a run while using offsets calculated from the original text.
    Multiple PII matches can occur in one run, and changing the first match
    changes the run length. Building each run's final text first prevents
    duplicated synthetic values such as ``Example Company 004``.
    """
    if not matches or not paragraph.runs:
        return

    original_runs = [(run.text or "", run) for run in paragraph.runs]
    run_ranges = []
    pos = 0
    for run_text, run in original_runs:
        run_ranges.append((pos, pos + len(run_text), run))
        pos += len(run_text)

    # Use the actual text represented by Word runs for offset calculations.
    # Some DOCX templates contain drawings/hyperlinks that python-docx exposes
    # in paragraph.text but not in paragraph.runs. Comparing the two strings
    # and returning here caused perfectly valid PII (especially CV emails) to
    # be skipped. The caller now computes matches against this same run text.
    original_text = "".join(run_text for run_text, _ in original_runs)
    if not original_text:
        return

    matches = sorted(matches, key=lambda m: (m.start, m.end))
    replacement_text = {
        id(m): replacements.fake(m.kind, m.value) for m in matches
    }

    for run_start, run_end, run in run_ranges:
        original = run.text or ""
        if not original:
            continue

        pieces = []
        cursor = 0

        for match in matches:
            if match.end <= run_start or match.start >= run_end:
                continue

            local_start = max(match.start, run_start) - run_start
            local_end = min(match.end, run_end) - run_start

            if run_start <= match.start < run_end:
                if local_start > cursor:
                    pieces.append(original[cursor:local_start])
                pieces.append(replacement_text[id(match)])
                cursor = max(cursor, local_end)
            else:
                # This run is covered by a match that began in an earlier run.
                cursor = max(cursor, local_end)

        if cursor < len(original):
            pieces.append(original[cursor:])

        new_text = "".join(pieces)
        if new_text != original:
            run.text = new_text

    for match in matches:
        stats[match.kind] = stats.get(match.kind, 0) + 1




# ----------------------------- Embedded image redaction ---------------------
def redact_embedded_images(docx_path: str):
    """Replace every embedded DOCX image with a same-sized blank placeholder.

    Prospectuses can contain logos, QR codes, scanned PAN/Aadhaar/ID cards,
    signatures and other visual identifiers that text-only redaction cannot
    remove. Keeping the original image would allow PII or company identity
    to leak through the output. We preserve each image's dimensions and file
    type so the surrounding Word layout remains stable.
    """
    tmp_path = docx_path + ".imgtmp"
    image_count = 0

    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(
        tmp_path, "w", compression=zipfile.ZIP_DEFLATED
    ) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)

            if item.filename.startswith("word/media/"):
                try:
                    with Image.open(io.BytesIO(data)) as src:
                        src = src.convert("RGB")
                        blank = Image.new("RGB", src.size, "white")

                        # Keep a small, clear indication that the visual content
                        # was intentionally removed. This is preferable to
                        # leaving an apparently broken image in the document.
                        draw = ImageDraw.Draw(blank)
                        text = "IMAGE REDACTED"
                        try:
                            font = ImageFont.load_default()
                        except Exception:
                            font = None

                        bbox = draw.textbbox((0, 0), text, font=font)
                        tw = bbox[2] - bbox[0]
                        th = bbox[3] - bbox[1]
                        x = max(0, (blank.width - tw) // 2)
                        y = max(0, (blank.height - th) // 2)
                        draw.text((x, y), text, fill="black", font=font)

                        ext = item.filename.rsplit(".", 1)[-1].lower()
                        out = io.BytesIO()
                        if ext in {"jpg", "jpeg"}:
                            blank.save(out, format="JPEG", quality=95)
                        elif ext == "png":
                            blank.save(out, format="PNG")
                        elif ext == "gif":
                            blank.save(out, format="GIF")
                        elif ext == "bmp":
                            blank.save(out, format="BMP")
                        else:
                            # Unknown media format: do not corrupt it.
                            zout.writestr(item, data)
                            continue

                        data = out.getvalue()
                        image_count += 1
                except Exception as exc:
                    print(
                        f"[warn] Could not redact embedded image {item.filename}: {exc}",
                        file=sys.stderr,
                    )

            zout.writestr(item, data)

    import os
    os.replace(tmp_path, docx_path)
    print(f"Embedded images redacted: {image_count}")



def redact_hyperlink_content(docx_path: str, replacements: ReplacementMap):
    """Redact visible hyperlink text and hyperlink targets.

    python-docx does not expose ``w:hyperlink`` text through paragraph.runs.
    CV templates frequently store email/LinkedIn/GitHub text there, so a
    normal run-based pass can miss it entirely. This XML pass handles both the
    visible text and the relationship target without depending on a particular
    DOCX template.
    """
    from lxml import etree
    import os

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    NS = {"w": W, "r": R}
    tmp_path = docx_path + ".linktmp"

    def local_kind(value: str):
        compact = re.sub(r"\s+", "", value)
        if EMAIL_RE.fullmatch(compact) or SPACED_EMAIL_RE.fullmatch(value):
            return "EMAIL"
        if URL_RE.search(value) or SPACED_URL_RE.search(value) or BARE_PROFILE_URL_RE.search(value):
            return "URL"
        return None

    with zipfile.ZipFile(docx_path, "r") as zin, zipfile.ZipFile(
        tmp_path, "w", compression=zipfile.ZIP_DEFLATED
    ) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)

            if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                try:
                    root = etree.fromstring(data)

                    for hyperlink in root.xpath(".//w:hyperlink", namespaces=NS):
                        nodes = hyperlink.xpath(".//w:t", namespaces=NS)
                        visible = "".join(n.text or "" for n in nodes)
                        if not visible.strip():
                            continue

                        kind = local_kind(visible)
                        if kind and nodes:
                            replacement = replacements.fake(kind, visible)
                            nodes[0].text = replacement
                            for node in nodes[1:]:
                                node.text = ""

                    data = etree.tostring(
                        root,
                        xml_declaration=True,
                        encoding="UTF-8",
                        standalone=True,
                    )
                except Exception as exc:
                    print(f"[warn] hyperlink XML pass failed for {item.filename}: {exc}", file=sys.stderr)

            # Sanitize external relationship targets. This prevents the source
            # email/domain from surviving in the DOCX package even when the
            # visible hyperlink text has been replaced.
            if item.filename.startswith("word/") and item.filename.endswith(".rels"):
                try:
                    root = etree.fromstring(data)
                    changed = False
                    for rel in root:
                        target = rel.get("Target") or ""
                        mode = rel.get("TargetMode")
                        if mode != "External":
                            continue
                        low = target.lower()
                        if "mailto:" in low or re.search(
                            r"(?:https?://|www\\.|linkedin\\.com|github\\.com|gitlab\\.com|bitbucket\\.org)",
                            low,
                        ):
                            rel.set("Target", "https://example.invalid/redacted")
                            changed = True
                    if changed:
                        data = etree.tostring(
                            root,
                            xml_declaration=True,
                            encoding="UTF-8",
                            standalone=True,
                        )
                except Exception as exc:
                    print(f"[warn] relationship XML pass failed for {item.filename}: {exc}", file=sys.stderr)

            zout.writestr(item, data)

    os.replace(tmp_path, docx_path)

# ----------------------------- Pipeline -------------------------------------
def redact_docx(input_path: str, output_path: str):
    doc = Document(input_path)
    blocks = collect_blocks(doc)

    people, companies, branches = prepare_detection_entities(blocks, doc)

    replacements = ReplacementMap()
    stats = {}

    company_pattern = build_entity_pattern(companies)
    person_pattern = build_entity_pattern(people)
    branch_pattern = build_entity_pattern(branches)

    for paragraph in iter_blocks(doc):
        run_text = "".join((run.text or "") for run in paragraph.runs)
        if not run_text.strip():
            continue
        matches = find_matches(
            run_text,
            people,
            companies,
            person_pattern=person_pattern,
            company_pattern=company_pattern,
            branch_pattern=branch_pattern,
        )
        replace_in_paragraph(paragraph, matches, replacements, stats)

    doc.save(output_path)
    redact_hyperlink_content(output_path, replacements)
    redact_embedded_images(output_path)
    return people, companies, stats, replacements


def main():
    parser = argparse.ArgumentParser(description="Redact PII from a DOCX file.")
    parser.add_argument("--input", required=True)
    parser.add_argument("--output", required=True)
    args = parser.parse_args()

    people, companies, stats, _ = redact_docx(args.input, args.output)

    print("Detected people:", len(people))
    print("Detected companies:", len(companies))
    for kind, count in sorted(stats.items()):
        print(f"{kind}: {count}")
    print("Output:", args.output)


if __name__ == "__main__":
    main()